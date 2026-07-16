from __future__ import annotations

import glob
import os
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Documentacion_Pruebas_E2E_Selenium_VargasVet.docx"
RESULT_IMAGE = Path(r"C:\Users\Usuario\AppData\Local\Temp\codex-clipboard-054ad622-d44a-4f49-b920-9e13676a48ee.png")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E8F5E9"
GREEN = "1B5E20"
PALE_AMBER = "FFF8E1"
AMBER = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
MUTED = "667085"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_cell_text(cell, size=8.5, color="000000", bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def set_cell_text(cell, text, size=8.5, color="000000", bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = str(text)
    style_cell_text(cell, size=size, color=color, bold=bold, alignment=alignment)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_section(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def setup_header_footer(section, right_label="VargasVet E2E | Selenium"):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = False
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("DOCUMENTACIÓN MAESTRA DE PRUEBAS E2E")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    r = p.add_run("  |  " + right_label)
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    add_page_field(fp)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc, text, size=28, color=NAVY, after=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    return p


def add_para(doc, text, size=11, bold=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT,
             after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    set_run_font(r, size=10, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=10, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_kv_table(doc, rows, widths=(2700, 6660)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, size=8.7, bold=True, color=NAVY)
        set_cell_shading(cells[0], LIGHT_GRAY)
        set_cell_text(cells[1], value, size=8.7)
    set_table_geometry(table, list(widths))
    return table


def add_steps_table(doc, steps):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["N.º", "Paso / acción", "Resultado esperado"]
    for i, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], text, size=8.5, bold=True, color=WHITE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for idx, (action, expected) in enumerate(steps, 1):
        cells = table.add_row().cells
        set_cell_text(cells[0], idx, size=8.4, bold=True, color=DARK_BLUE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], action, size=8.4)
        set_cell_text(cells[2], expected, size=8.4)
    set_table_geometry(table, [600, 4140, 4620])
    return table


def result_durations():
    result = {}
    for file in (ROOT / "target" / "surefire-reports").glob("TEST-*.xml"):
        tree = ET.parse(file)
        for tc in tree.getroot().findall("testcase"):
            result[tc.attrib["name"]] = float(tc.attrib.get("time", "0"))
    return result


def latest_evidence(case_number: str, extension: str):
    folder = ROOT / "target" / "evidence" / ("videos" if extension == "mp4" else "screenshots")
    pattern = str(folder / f"CASO-{case_number}-*.{extension}")
    files = sorted(glob.glob(pattern), key=os.path.getmtime, reverse=True)
    return Path(files[0]).name if files else "No disponible"


CASES = [
    dict(id="CP-E2E-001", num="01", method="caso01_reserva_concurrente_del_mismo_horario",
         title="Solo una sesión reserva el mismo horario concurrente", module="Agenda de citas",
         function="Reserva concurrente y exclusión de horario", priority="Alta", type="Negocio, concurrencia e integración",
         actor="Administrador (dos sesiones simultáneas)",
         objective="Comprobar que dos solicitudes simultáneas para el mismo veterinario, fecha y hora no generen una doble reserva.",
         pre="Backend y frontend E2E activos; base H2 inicializada; administrador, apoderado, mascota, servicio y veterinario habilitados; horario 14:00 disponible.",
         data="Fecha de concurrencia del fixture; 14:00; motivo E2E CONCURRENCIA; dos WebDriver independientes.",
         deps="Autenticación, API de citas, bloqueo transaccional, repositorios JPA y persistencia H2.",
         steps=[("Iniciar sesión como administrador en dos navegadores.", "Ambas sesiones acceden a la agenda con permisos de escritura."),
                ("Preparar en ambas sesiones la misma cita: dueño, mascota, servicio, veterinario, fecha y 14:00.", "Los formularios quedan válidos con datos equivalentes."),
                ("Enviar las dos solicitudes en paralelo.", "El backend serializa la operación crítica."),
                ("Comparar los dos resultados.", "Una sesión recibe 'Cita programada' y la otra un conflicto de horario."),
                ("Verificar la regla final.", "Existe exactamente una reserva; no hay duplicidad en base de datos.")],
         final="Una sola cita queda registrada para el intervalo; la solicitud competidora es rechazada con explicación de conflicto.",
         acceptance="1 éxito + 1 rechazo; cero citas duplicadas para veterinario/intervalo.",
         controls="Frontend: resultado visible por sesión. Backend: bloqueo pesimista y validación de solapamiento. BD: unicidad lógica del intervalo. Manejo de error: conflicto controlado. Logs/auditoría: registrar creación aceptada y rechazo técnico sin stack trace al usuario."),
    dict(id="CP-E2E-002", num="02", method="caso02_disponibilidad_integral",
         title="Disponibilidad según clínica, turno, duración y ocupación", module="Agenda de citas",
         function="Cálculo integral de horarios disponibles", priority="Alta", type="Negocio e integración",
         actor="Administrador", objective="Validar que la agenda solo ofrezca intervalos operativos y libres.",
         pre="Clínica abierta de 08:00 a 18:00; veterinario con turno activo; citas preexistentes a las 09:00, 10:00 y 12:00.",
         data="Ana Pruebas; Luna E2E; Consulta general E2E; Victor Veterinario; duración 30 minutos.",
         deps="Horarios de empresa, turnos del empleado, duración del servicio y citas persistidas.",
         steps=[("Autenticarse como administrador y abrir Agenda.", "Se muestra la vista de agenda."),
                ("Abrir Nueva Cita y seleccionar dueño, mascota, servicio y veterinario.", "El formulario acepta las relaciones válidas."),
                ("Seleccionar la fecha preparada.", "Se calcula la lista de intervalos disponibles."),
                ("Revisar 09:00, 10:00 y 12:00.", "Los horarios ocupados no se ofrecen."),
                ("Revisar todos los intervalos restantes.", "Cada intervalo está entre 08:00 y 18:00 y admite los 30 minutos completos.")],
         final="La disponibilidad presentada coincide con las restricciones combinadas de operación, turno, duración y ocupación.",
         acceptance="Lista no vacía; ocupados excluidos; 100% de opciones dentro del horario de la clínica.",
         controls="Frontend: listas dependientes y horarios. Backend: cálculo de disponibilidad. BD: consulta de citas/turnos. Errores: lista vacía controlada. Auditoría: no aplica por ser consulta; logs de error solo ante fallo de API."),
    dict(id="CP-E2E-003", num="03", method="caso03_limites_diarios_del_apoderado",
         title="Límites diarios por mascota y apoderado", module="Portal del apoderado",
         function="Cuotas de citas diarias", priority="Alta", type="Negocio, negativo e integración",
         actor="Apoderado", objective="Comprobar el máximo de 2 citas por mascota y 3 citas totales por apoderado en un día.",
         pre="Apoderado autenticable con tres mascotas activas; servicio y veterinario disponibles en la fecha de cuota.",
         data="Luna, Nube y Sol E2E; horas 09:00, 10:00, 11:00 y 12:00; motivos CUOTA UNO/DOS/TRES/CUATRO.",
         deps="Portal, API de citas, reglas de cuota y persistencia.",
         steps=[("Iniciar sesión como apoderado.", "Se abre el portal autorizado."),
                ("Crear dos citas para Luna en la misma fecha.", "Ambas citas se registran."),
                ("Intentar una tercera cita para Luna.", "Se rechaza con mensaje 'más de 2 citas'."),
                ("Crear la tercera cita total usando Nube.", "La cita se registra porque no excede el total diario."),
                ("Intentar una cuarta cita total usando Sol.", "Se rechaza con mensaje 'más de 3 citas'.")],
         final="Se respetan simultáneamente las cuotas por mascota y por apoderado.",
         acceptance="La tercera de la misma mascota y la cuarta total son rechazadas; las tres permitidas persisten.",
         controls="Frontend: mensajes y refresco de portal. Backend: conteos diarios por mascota/apoderado. BD: solo tres altas válidas. Seguridad: el apoderado usa solo sus mascotas. Auditoría/logs: altas aceptadas y rechazos de negocio trazables."),
    dict(id="CP-E2E-004", num="04", method="caso04_inicio_controlado_de_atencion",
         title="Inicio y cierre por veterinario asignado", module="Atención clínica",
         function="Ciclo de consulta clínica", priority="Alta", type="Negocio, permisos e integración",
         actor="Veterinario asignado", objective="Validar que el veterinario asignado inicia, completa datos obligatorios y cierra la atención.",
         pre="Cita E2E CONSULTA programada y restaurada; veterinario asignado activo con permiso clínico.",
         data="Tipo Control rutina; peso 12.8; anamnesis, examen físico y observaciones E2E.",
         deps="Agenda, historias clínicas, autorización, autoguardado, control de versión y estados de cita/consulta.",
         steps=[("Autenticarse como veterinario y abrir la cita asignada.", "La acción Iniciar consulta está disponible."),
                ("Iniciar la consulta.", "Se navega a /historias-clinicas/consulta/{id}."),
                ("Completar tipo, peso, anamnesis, examen físico y observaciones.", "El formulario queda válido y se autoguarda."),
                ("Esperar el incremento de versión.", "La versión cambia, confirmando persistencia."),
                ("Cerrar y confirmar la consulta.", "La consulta se cierra y se vuelve a Agenda.")],
         final="Consulta cerrada con datos obligatorios persistidos y cita completada.",
         acceptance="Ruta clínica válida, versión incrementada y cierre confirmado sin pérdida de información.",
         controls="Frontend: validación obligatoria y autoguardado. Backend: permiso del veterinario asignado y transición de estados. BD: Consulta CERRADA/Cita COMPLETADA. Auditoría: ACTUALIZAR_CONSULTA y CERRAR_CONSULTA. Logs: ausencia de errores no controlados."),
    dict(id="CP-E2E-005", num="13", method="caso13_desactivacion_de_empresa",
         title="Empresa inactiva bloquea usuarios y puede reactivarse", module="Administración y autenticación",
         function="Ciclo de estado de empresa", priority="Alta", type="Seguridad, permisos y negocio",
         actor="Superadministrador y administrador de empresa", objective="Verificar que una empresa inactiva impida el acceso de sus usuarios y que la reactivación restablezca el inicio de sesión.",
         pre="Clínica Desactivable E2E activa; usuario administrador asociado y credenciales válidas.",
         data="Empresa Clínica Desactivable E2E; usuario e2e.company.admin@vargasvet.test.",
         deps="Administración de empresas, autenticación JWT, estado de usuario/empresa y persistencia.",
         steps=[("Ingresar como superadministrador y desactivar la empresa.", "El estado queda Inactivo."),
                ("En otra sesión, intentar iniciar con el administrador afectado.", "El login es rechazado con mensaje relacionado con empresa."),
                ("Reactivar la empresa en un bloque de restauración.", "El estado vuelve a Activo incluso ante fallo intermedio."),
                ("Iniciar sesión nuevamente con el usuario afectado.", "El acceso se concede y se abandona /login.")],
         final="El estado de la empresa gobierna el acceso de sus usuarios y es reversible.",
         acceptance="Bloqueo mientras está inactiva y acceso exitoso después de reactivación.",
         controls="Frontend: confirmación de cambio de estado y mensaje de login. Backend: validación empresa-activa al autenticar. BD: estado restaurado. Seguridad: no emitir sesión/JWT durante bloqueo. Auditoría: desactivación/reactivación e intentos de acceso."),
    dict(id="CP-E2E-006", num="14", method="caso14_desactivacion_de_rol",
         title="Rol inactivo deja de autorizar al usuario", module="Roles y autenticación",
         function="Revocación por estado de rol", priority="Alta", type="Seguridad y permisos",
         actor="Superadministrador y usuario restringido", objective="Validar que un usuario sin rol activo no pueda autenticarse ni conservar autorización efectiva.",
         pre="Rol ROLE_E2E_RESTRICTED activo y asignado a un usuario de prueba.",
         data="Usuario e2e.restricted@vargasvet.test; rol restringido E2E.",
         deps="Gestión de roles, autenticación, resolución de authorities y base de datos.",
         steps=[("Ingresar como superadministrador y desactivar el rol.", "El rol queda Inactivo."),
                ("Intentar iniciar sesión con el usuario asignado.", "El acceso se rechaza por falta de rol activo."),
                ("Restaurar el rol en el bloque de limpieza.", "El rol vuelve a Activo.")],
         final="La desactivación del rol revoca la autorización del usuario asociado.",
         acceptance="Login rechazado con referencia a rol activo; estado restaurado al finalizar.",
         controls="Frontend: cambio de estado y error visible. Backend: filtrar roles inactivos al autenticar. Seguridad: no emitir JWT utilizable. BD: relación usuario-rol conservada, rol reactivado. Auditoría/logs: cambio de rol e intento rechazado."),
    dict(id="CP-E2E-007", num="16", method="caso16_reprogramacion_con_reglas",
         title="Reprogramación conserva datos y disponibilidad", module="Agenda de citas",
         function="Reprogramación", priority="Alta", type="Negocio e integración",
         actor="Administrador", objective="Verificar que una cita sea reprogramada a un horario válido sin perder su identidad ni motivo.",
         pre="Cita E2E REPROGRAMAR existente; nueva fecha con horario 14:00 disponible.",
         data="Motivo E2E REPROGRAMAR; fecha objetivo del fixture; 14:00.",
         deps="Agenda, disponibilidad, control de versión, actualización de cita y persistencia.",
         steps=[("Iniciar sesión y localizar E2E REPROGRAMAR.", "La cita aparece en la fecha original."),
                ("Abrir acciones y seleccionar Reprogramar.", "Se abre el formulario con los datos existentes."),
                ("Elegir nueva fecha y 14:00.", "El horario válido puede seleccionarse."),
                ("Enviar y confirmar 'Sí, reprogramar'.", "La API actualiza la cita."),
                ("Abrir la fecha destino.", "La cita aparece con estado Reprogramada y el mismo motivo.")],
         final="La cita queda en la nueva fecha/hora y conserva sus datos funcionales.",
         acceptance="Registro visible en destino con motivo original y estado Reprogramada.",
         controls="Frontend: formulario precargado y confirmación. Backend: disponibilidad y actualización versionada. BD: mismo registro lógico, nueva fecha/hora. Auditoría: reprogramación. Manejo de error: impedir horario ocupado."),
    dict(id="CP-E2E-008", num="18", method="caso18_emergencia_fuera_del_horario",
         title="Emergencia autorizada con clínica cerrada", module="Agenda y emergencias",
         function="Excepción operativa por emergencia", priority="Alta", type="Negocio, permisos e integración",
         actor="Administrador autorizado", objective="Comprobar que una emergencia permitida pueda agendarse cuando la clínica está cerrada.",
         pre="Escenario E2E marca la clínica cerrada hoy; servicio permite emergencia; actor autorizado.",
         data="Luna E2E; Consulta general E2E; motivo EMERGENCIA SELENIUM; fecha actual.",
         deps="Horario de clínica, atributo permiteEmergencia, autorización y API de citas.",
         steps=[("Cerrar la clínica mediante fixture controlado.", "El horario operativo de hoy queda cerrado."),
                ("Autenticarse y completar una nueva cita.", "El formulario detecta el cierre."),
                ("Seleccionar Agendar como Emergencia o activar la marca equivalente.", "Se habilita la excepción autorizada."),
                ("Registrar el motivo y enviar.", "La cita se programa como emergencia."),
                ("Restaurar el horario de la clínica.", "El entorno vuelve a abierto aunque falle el caso.")],
         final="La emergencia se registra pese al cierre sin alterar permanentemente el horario.",
         acceptance="Mensaje de programación/emergencia y restauración garantizada del escenario.",
         controls="Frontend: opción de emergencia contextual. Backend: validar autorización y servicio compatible. BD: esEmergencia=true. Auditoría: creación excepcional. Logs: motivo del override; error controlado si actor no autorizado."),
    dict(id="CP-E2E-009", num="21", method="caso21_desactivacion_de_empleado",
         title="Empleado inactivo desaparece de disponibilidad", module="Personal y agenda",
         function="Disponibilidad por estado del empleado", priority="Alta", type="Negocio e integración",
         actor="Administrador", objective="Validar que un empleado inactivo no sea seleccionable para nuevas citas.",
         pre="Empleado E2E activo, con servicio y horario asignados.",
         data="Empleado VargasVet E2E de alternancia; Luna E2E; Consulta general E2E.",
         deps="Gestión de personal, disponibilidad de agenda y filtros backend.",
         steps=[("Desactivar al empleado desde Personal.", "El estado cambia a Inactivo."),
                ("Abrir Nueva Cita y seleccionar dueño, mascota y servicio.", "Se habilita la selección de veterinario."),
                ("Abrir la lista de empleados/veterinarios.", "El empleado inactivo no aparece."),
                ("Reactivar al empleado en la limpieza.", "El entorno queda restaurado.")],
         final="Ninguna nueva cita puede asignarse al empleado inactivo.",
         acceptance="Ausencia del empleado en la lista y reactivación posterior.",
         controls="Frontend: lista filtrada. Backend: excluir estado=false. BD: estado reversible; citas históricas intactas. Auditoría: cambio de estado. Manejo de error: rechazar payload manipulado con empleado inactivo (pendiente de prueba API)."),
    dict(id="CP-E2E-010", num="24", method="caso24_aislamiento_de_datos_del_apoderado",
         title="Aislamiento de historial entre apoderados", module="Portal e historia clínica",
         function="Autorización por pertenencia de mascota", priority="Alta", type="Seguridad, negativo e IDOR",
         actor="Apoderado", objective="Comprobar que manipular el identificador de mascota no exponga el historial de otro apoderado.",
         pre="Ana Pruebas tiene Luna, Nube y Sol; otro apoderado posee Mascota Ajena E2E.",
         data="petId autorizado y foreignPetId ajeno enviados directamente en URL.",
         deps="Autenticación, portal, endpoint de mascotas/historia y autorización por propietario.",
         steps=[("Iniciar sesión como Ana Pruebas.", "El portal carga únicamente su contexto."),
                ("Abrir /mi-historial/{petId autorizado}.", "Se muestra Luna E2E."),
                ("Manipular la URL con foreignPetId.", "El portal conserva la lista de mascotas autorizadas."),
                ("Inspeccionar el contenido visible.", "Mascota Ajena E2E y su historial no aparecen.")],
         final="La manipulación de ID no revela datos ajenos y el usuario permanece en su ámbito autorizado.",
         acceptance="Mascota autorizada visible; nombre de mascota ajena ausente.",
         controls="Frontend: no preseleccionar ID no autorizado. Backend: ownership obligatorio, no confiar en filtros UI. Seguridad: prevención IDOR. BD: ninguna modificación. Auditoría/logs: registrar intento denegado sin revelar existencia del recurso."),
    dict(id="CP-E2E-011", num="27", method="caso27_control_de_version_en_consulta",
         title="Edición clínica concurrente con control de versión", module="Atención clínica",
         function="Optimistic locking de consulta", priority="Alta", type="Concurrencia, integridad y negativo",
         actor="Veterinario en dos sesiones", objective="Impedir que una edición obsoleta sobrescriba una modificación clínica confirmada.",
         pre="Consulta concurrente restaurada y abierta por el veterinario en dos navegadores.",
         data="Anamnesis de primera sesión y edición obsoleta diferenciada; token de versión.",
         deps="Autoguardado Angular, API de consulta, campo version y transacción JPA.",
         steps=[("Abrir la misma consulta en dos sesiones autenticadas.", "Ambas cargan la versión inicial."),
                ("Editar datos obligatorios en la primera sesión.", "El autoguardado persiste e incrementa la versión."),
                ("Editar con la segunda sesión obsoleta.", "El autoguardado entra en Error al guardar."),
                ("Pulsar Reintentar.", "Se muestra 'modificada por otro usuario'."),
                ("Cerrar la consulta desde la primera sesión.", "Se conserva la edición válida y se limpia el escenario.")],
         final="La segunda edición no sobrescribe datos; el conflicto se comunica al usuario.",
         acceptance="Versión cambia tras primera edición y la segunda recibe conflicto explícito.",
         controls="Frontend: estado de autoguardado y reintento. Backend: comparación request.version vs entidad. BD: integridad de la última versión válida. Auditoría: solo actualización aceptada. Logs: conflicto esperado, sin 500 genérico."),
    dict(id="CP-E2E-012", num="28", method="caso28_desactivacion_de_servicio",
         title="Servicio inactivo no admite nuevas citas y conserva historial", module="Servicios y agenda",
         function="Ciclo de servicio e integridad histórica", priority="Alta", type="Negocio, integración y regresión",
         actor="Administrador", objective="Validar que un servicio inactivo deje de ofrecerse sin borrar citas históricas.",
         pre="Servicio alternable activo y cita histórica E2E SERVICIO HISTORICO existente.",
         data="Servicio E2E alternable; Nube E2E; fecha histórica y fecha para nueva cita.",
         deps="Configuración complementaria, agenda, filtros de servicios y persistencia histórica.",
         steps=[("Desactivar el servicio.", "El servicio queda Inactivo."),
                ("Abrir Nueva Cita y lista de servicios.", "El servicio inactivo no aparece para nuevas operaciones."),
                ("Abrir la agenda de la fecha histórica.", "La cita E2E SERVICIO HISTORICO sigue visible."),
                ("Reactivar el servicio.", "El estado se restaura para otros casos.")],
         final="Se bloquea el uso futuro sin pérdida de trazabilidad histórica.",
         acceptance="Servicio ausente en altas; cita histórica presente; reactivación exitosa.",
         controls="Frontend: filtro en selector y visualización histórica. Backend: rechazar servicio inactivo en nuevas citas. BD: no eliminar referencias históricas. Auditoría: cambio de estado. Integridad referencial preservada."),
    dict(id="CP-E2E-013", num="31", method="caso31_activacion_de_cuenta",
         title="Activación de cuenta e inicio de sesión", module="Autenticación y empleados",
         function="Verificación inicial y establecimiento de contraseña", priority="Alta", type="Seguridad y flujo positivo",
         actor="Empleado nuevo", objective="Comprobar que un token válido permita establecer contraseña, activar la cuenta e iniciar sesión.",
         pre="Empleado Alicia Activación inactivo, email no verificado, passwordChanged=false y token E2E válido.",
         data="Token selenium-e2e-activation-token; contraseña nueva Activated!456 y confirmación idéntica.",
         deps="Endpoint de verificación, política de contraseña, persistencia de usuario y autenticación JWT.",
         steps=[("Restablecer el escenario de activación.", "Cuenta y token vuelven al estado inicial."),
                ("Abrir /auth/verify/{token}.", "Se presenta el formulario de nueva contraseña."),
                ("Ingresar y confirmar Activated!456.", "Los campos cumplen la política y coinciden."),
                ("Enviar el formulario.", "Se muestra 'Cuenta activada'."),
                ("Iniciar sesión con email y nueva contraseña.", "El usuario sale de /login y obtiene sesión válida.")],
         final="La cuenta queda activa, verificada y accesible con la nueva credencial.",
         acceptance="Mensaje de activación y login exitoso posterior.",
         controls="Frontend: coincidencia y fortaleza de contraseña. Backend: token válido/no reutilizable, hash seguro y activación atómica. BD: activo/emailVerified/passwordChanged actualizados. Auditoría: activación y login. Logs: token inválido/expirado sin exponer secretos (pendiente negativo)."),
    dict(id="CP-E2E-014", num="40", method="caso40_flujo_clinico_con_radiografia_e_ia",
         title="Historia, radiografía y asistencia IA", module="Atención clínica e IA",
         function="Flujo clínico integrado con almacenamiento y análisis", priority="Alta", type="Negocio, integración y archivo",
         actor="Veterinario", objective="Validar el flujo de consulta, carga de radiografía, cierre e interpretación integrada mediante IA.",
         pre="Cita E2E CONSULTA IA restaurada; veterinario asignado; almacenamiento y respuesta IA E2E disponibles.",
         data="Sol E2E; radiografia-selenium.png; tipo Radiografía; descripción; anamnesis clínica.",
         deps="Agenda, consulta, autoguardado, upload multipart, almacenamiento, historia clínica y servicio IA.",
         steps=[("Iniciar la consulta IA y completar datos obligatorios.", "La consulta se autoguarda e incrementa versión."),
                ("Abrir Exámenes y adjuntar radiografia-selenium.png.", "El input oculto recibe el archivo y habilita la carga."),
                ("Seleccionar tipo Radiografía, describir y subir.", "El archivo aparece asociado a la consulta."),
                ("Cerrar la consulta.", "Historia y archivo quedan consolidados."),
                ("Abrir la historia de Sol E2E y solicitar análisis IA.", "Se muestra análisis 'HC + Radiografía' y resultado 'sin hallazgos óseos agudos'.")],
         final="El dato clínico, el archivo y el análisis IA quedan integrados de extremo a extremo.",
         acceptance="Archivo visible, consulta cerrada y respuesta IA esperada presentada.",
         controls="Frontend: validación de formulario/archivo y estados de carga. Backend: autorización, multipart y asociación a consulta. Almacenamiento: archivo recuperable. BD: metadatos íntegros. IA: manejo de respuesta. Auditoría/logs: subida, cierre y solicitud IA; timeout/error externo aún pendiente."),
]


BACKLOG = [
    ("CP-E2E-P001", "Autenticación", "Token de activación expirado, inválido o reutilizado", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P002", "Seguridad", "JWT ausente, vencido o alterado en acceso directo por URL", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P003", "Seguridad", "Manipulación de payload para companyId, ownerId, empleado o servicio inactivo", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P004", "Entradas", "SQL Injection y XSS en búsquedas, motivos y campos clínicos", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P005", "Archivos", "Extensión inválida y archivo superior a 20 MB", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P006", "Recursos", "ID inexistente, eliminado o perteneciente a otra empresa", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P007", "Formularios", "Campos obligatorios vacíos, longitudes y caracteres no permitidos", "Media", "Pendiente / Supuesto"),
    ("CP-E2E-P008", "Integraciones", "Timeout/error 5xx del servicio IA y recuperación del usuario", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P009", "Navegación", "Búsquedas, filtros, orden y paginación con volúmenes altos", "Media", "Pendiente / Supuesto"),
    ("CP-E2E-P010", "Exportación", "Exportación PDF y consistencia con los datos clínicos", "Media", "Pendiente / Supuesto"),
    ("CP-E2E-P011", "Seguridad web", "CORS y CSRF según arquitectura de autenticación", "Alta", "Pendiente / Supuesto"),
    ("CP-E2E-P012", "Rendimiento", "Carga de agenda, concurrencia sostenida, memoria y tiempos de API", "Alta", "Pendiente / Supuesto"),
]


TRACE = [
    ("RN-E2E-01", "Evitar doble reserva y solapamientos", "CP-E2E-001, CP-E2E-002", "Cubierto"),
    ("RN-E2E-02", "Aplicar cuotas diarias del apoderado", "CP-E2E-003", "Cubierto"),
    ("RN-E2E-03", "Controlar inicio/cierre de atención", "CP-E2E-004", "Cubierto"),
    ("RN-E2E-04", "Bloquear acceso por empresa o rol inactivo", "CP-E2E-005, CP-E2E-006", "Cubierto"),
    ("RN-E2E-05", "Reprogramar conservando identidad y reglas", "CP-E2E-007", "Cubierto"),
    ("RN-E2E-06", "Permitir emergencia autorizada", "CP-E2E-008", "Cubierto"),
    ("RN-E2E-07", "Excluir empleados inactivos", "CP-E2E-009", "Cubierto"),
    ("RN-E2E-08", "Aislar datos por apoderado", "CP-E2E-010", "Cubierto"),
    ("RN-E2E-09", "Evitar sobrescritura clínica concurrente", "CP-E2E-011", "Cubierto"),
    ("RN-E2E-10", "Inactivar servicio sin perder historial", "CP-E2E-012", "Cubierto"),
    ("RN-E2E-11", "Activar cuenta con token válido", "CP-E2E-013", "Cubierto"),
    ("RN-E2E-12", "Integrar historia, radiografía e IA", "CP-E2E-014", "Cubierto"),
]


def build_document():
    durations = result_durations()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], landscape=False)
    setup_header_footer(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True

    # Cover: editorial_cover override over compact_reference_guide.
    add_para(doc, "VARGASVET E2E", size=11, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=54)
    add_title(doc, "Documentación maestra de pruebas End-to-End", size=27, color=NAVY, after=8)
    add_para(doc, "Selenium WebDriver · Java · JUnit 5", size=14, color=DARK_BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "Plan, diseño, ejecución, trazabilidad y cierre de pruebas",
             size=11.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
    add_kv_table(doc, [
        ("Sistema", "VargasVet - Sistema de gestión veterinaria"),
        ("Documento", "Especificación e informe maestro de pruebas E2E"),
        ("Versión", "1.0"),
        ("Fecha de ejecución", "14 de julio de 2026"),
        ("Responsable", "Equipo de Aseguramiento de Calidad (QA)"),
        ("Marco de referencia", "Buenas prácticas ISTQB e ISO/IEC/IEEE 29119"),
        ("Resultado", "14 ejecutados · 14 aprobados · 0 fallos · 0 errores · 0 omitidos"),
    ])
    add_para(doc, "Estado del documento: APROBADO", size=12, bold=True, color=GREEN,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Clasificación: Uso académico / técnico", size=9.5, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    doc.add_page_break()
    doc.add_heading("Control documental", level=1)
    add_kv_table(doc, [
        ("Propósito", "Definir y evidenciar los casos E2E automatizados que validan reglas críticas de negocio, seguridad e integración."),
        ("Alcance", "Los 14 casos seleccionados e implementados en selenium-e2e. No representa cobertura total de todas las funcionalidades del producto."),
        ("Fuentes", "Código Selenium, fixtures E2E, reportes JUnit/Surefire, capturas, videos y salida final BUILD SUCCESS."),
        ("Convención de supuestos", "Toda inferencia no confirmada por requisito formal se marca como 'Supuesto'."),
        ("Criterio de estado", "Aprobado = ejecución automatizada sin failure/error. Pendiente = sugerido, aún no implementado o no ejecutado."),
    ])
    doc.add_heading("Historial de versiones", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("Versión", "Fecha", "Cambio", "Responsable")):
        set_cell_text(table.rows[0].cells[i], h, size=8.5, bold=True, color=WHITE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    row = table.add_row().cells
    for i, v in enumerate(("1.0", "14/07/2026", "Emisión inicial con ejecución 14/14 aprobada", "QA")):
        set_cell_text(row[i], v, size=8.5)
    set_table_geometry(table, [1200, 1500, 4960, 1700])

    doc.add_heading("Contenido", level=1)
    contents = [
        "1. Resumen ejecutivo", "2. Estrategia, alcance y ambiente", "3. Resultado de ejecución",
        "4. Catálogo resumido", "5. Especificación detallada de 14 casos", "6. Trazabilidad y cobertura",
        "7. Riesgos, defectos y recomendaciones", "8. Backlog de pruebas pendientes", "9. Evidencias y cierre",
    ]
    for item in contents:
        add_para(doc, item, size=10.5, color=DARK_BLUE, after=3)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    add_para(doc, "La suite E2E de VargasVet fue ejecutada el 14/07/2026 en Windows 11 con Chrome, Selenium WebDriver 4.33.0, Java 22, JUnit 5, frontend Angular en http://127.0.0.1:4200 y backend Spring Boot en http://127.0.0.1:8080/api/v1. La ejecución final concluyó con BUILD SUCCESS: 14 pruebas ejecutadas, 14 aprobadas, 0 fallos, 0 errores y 0 omitidas.")
    add_callout(doc, "Conclusión QA", "El alcance crítico seleccionado es apto para cierre de la iteración E2E. La evidencia respalda las reglas de agenda, concurrencia clínica, ciclo de entidades, aislamiento de datos, activación de cuenta e integración con archivos/IA.", fill=PALE_GREEN, color=GREEN)
    add_callout(doc, "Limitación", "El 100% corresponde al alcance seleccionado (14 casos), no a todas las funcionalidades posibles del sistema. Las brechas identificadas se registran como backlog pendiente.", fill=PALE_AMBER, color=AMBER)

    doc.add_heading("Indicadores", level=2)
    metrics = doc.add_table(rows=2, cols=4)
    metrics.style = "Table Grid"
    metric_data = [("Ejecución", "14/14"), ("Aprobación", "100%"), ("Automatización", "100%"), ("Defectos abiertos", "0")]
    for i, (label, value) in enumerate(metric_data):
        set_cell_text(metrics.rows[0].cells[i], value, size=17, bold=True, color=GREEN, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(metrics.rows[0].cells[i], PALE_GREEN)
        set_cell_text(metrics.rows[1].cells[i], label, size=8.5, bold=True, color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(metrics, [2340, 2340, 2340, 2340])

    doc.add_heading("2. Estrategia, alcance y ambiente", level=1)
    add_kv_table(doc, [
        ("Enfoque", "Pruebas E2E basadas en riesgo y reglas de negocio; caja negra a través de UI con verificación de efectos integrados."),
        ("Niveles integrados", "Frontend ↔ Backend ↔ H2; autenticación; almacenamiento; servicio de IA simulado/controlado en perfil E2E."),
        ("Diseño", "Page Object Model, fixtures deterministas, sesiones paralelas donde aplica y restauración de estados mediante finally."),
        ("Entrada", "Backend y frontend disponibles, puertos 8080/4200 libres al inicio, dependencias instaladas, Chrome y FFmpeg en PATH."),
        ("Salida", "Todos los casos terminan sin fallos/errores; evidencia generada; procesos iniciados por el runner finalizados."),
        ("No incluido", "Carga masiva, compatibilidad multinavegador, pentesting, CORS/CSRF, caos de red y fallos reales de servicios externos."),
    ])
    doc.add_heading("Ambiente de ejecución", level=2)
    add_kv_table(doc, [
        ("Sistema operativo", "Windows 11, arquitectura amd64"),
        ("Navegador", "Google Chrome 150.0.7871.115 (según reporte Selenium)"),
        ("Automatización", "Selenium 4.33.0 + JUnit Jupiter 5.12.2 + Maven Surefire 3.5.3"),
        ("Java", "22.0.2; compilación del módulo con release 21"),
        ("Datos", "Perfil e2e con H2 en memoria y fixtures aislados"),
        ("Evidencia", "Captura final por caso, video MP4 por caso, XML JUnit y logs de servidores"),
        ("Comando", ".\\run-e2e.ps1 -Video -SlowMo 500"),
    ])
    doc.add_heading("Criterios de aceptación generales", level=2)
    add_para(doc, "Cada caso debe completar su flujo sin excepción no controlada, demostrar la regla mediante una aserción observable, conservar la integridad de datos y generar evidencia. Los escenarios que mutan configuraciones deben restaurar su estado aunque fallen.")

    doc.add_heading("3. Resultado de ejecución", level=1)
    if RESULT_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(RESULT_IMAGE), width=Inches(6.35))
        picture._inline.docPr.set("descr", "Salida final de Maven: 14 pruebas ejecutadas y aprobadas, BUILD SUCCESS")
        picture._inline.docPr.set("title", "Resultado final de la suite Selenium E2E")
        add_para(doc, "Figura 1. Salida final de Maven/Surefire: Tests run 14, Failures 0, Errors 0, Skipped 0; BUILD SUCCESS.", size=8.5, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_kv_table(doc, [
        ("Fecha/hora de cierre", "14/07/2026 15:08:35 (America/Lima)"),
        ("Duración total", "03:53 min (salida final proporcionada)"),
        ("Resultado", "APROBADO"),
        ("Tasa de aprobación", "100% (14/14)"),
        ("Fallos / errores / omitidos", "0 / 0 / 0"),
        ("Advertencia no bloqueante", "Selenium informó ausencia de implementación CDP exacta para Chrome 150; no afectó WebDriver ni el resultado."),
    ])

    # Landscape catalog summary (named geometry override).
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    setup_header_footer(section, "Catálogo resumido")
    doc.add_heading("4. Catálogo resumido de casos", level=1)
    add_para(doc, "Tabla solicitada de diseño y resultado. Los pasos y resultados completos se desarrollan en la sección 5.", size=9.5, color=MUTED)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["ID", "Módulo", "Objetivo", "Precondiciones", "Datos", "Pasos", "Resultado esperado", "Prioridad", "Estado"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, size=7.2, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for case in CASES:
        cells = table.add_row().cells
        values = [case["id"], case["module"], case["objective"], case["pre"], case["data"],
                  f"{len(case['steps'])} pasos", case["final"], case["priority"], "Aprobado"]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, size=6.8, bold=(i in (0, 7, 8)),
                          color=(GREEN if i == 8 else "000000"))
        set_cell_shading(cells[8], PALE_GREEN)
    set_table_geometry(table, [1050, 1350, 2050, 1900, 1650, 700, 2000, 750, 850], indent_dxa=120)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    setup_header_footer(section, "Especificación detallada")
    doc.add_heading("5. Especificación detallada de casos", level=1)
    add_para(doc, "Cada ficha contiene diseño, ejecución y controles técnicos. Los campos de auditoría/logs indicados como 'esperados' son criterios de aseguramiento; la suite actual no inspecciona todos ellos de forma directa.")

    for idx, case in enumerate(CASES):
        if idx > 0:
            doc.add_page_break()
        doc.add_heading(f"{case['id']} · CASO {case['num']}", level=1)
        add_para(doc, case["title"], size=14, bold=True, color=NAVY, after=8)
        duration = durations.get(case["method"], 0.0)
        add_kv_table(doc, [
            ("Módulo", case["module"]),
            ("Funcionalidad", case["function"]),
            ("Objetivo", case["objective"]),
            ("Prioridad", case["priority"]),
            ("Tipo de prueba", case["type"]),
            ("Actor", case["actor"]),
            ("Precondiciones", case["pre"]),
            ("Datos de prueba", case["data"]),
            ("Dependencias", case["deps"]),
            ("Estado", f"APROBADO · ejecución {duration:.3f} s"),
        ])
        doc.add_heading("Flujo completo y resultados por paso", level=2)
        add_steps_table(doc, case["steps"])
        doc.add_heading("Oráculo y criterio de aceptación", level=2)
        add_kv_table(doc, [
            ("Resultado esperado final", case["final"]),
            ("Criterio de aceptación", case["acceptance"]),
            ("Validaciones y controles", case["controls"]),
            ("Mensajes esperados", "Los descritos en los resultados por paso; cualquier excepción técnica debe transformarse en mensaje comprensible y no exponer información sensible."),
            ("Evidencias requeridas", f"Captura: {latest_evidence(case['num'], 'png')} | Video: {latest_evidence(case['num'], 'mp4')} | XML JUnit y logs de servidor de la ejecución."),
        ])

    doc.add_page_break()
    doc.add_heading("6. Trazabilidad y cobertura", level=1)
    doc.add_heading("Matriz Requisito → Caso de prueba", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("Requisito", "Regla / capacidad", "Casos", "Estado")):
        set_cell_text(table.rows[0].cells[i], h, size=8.5, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for req, desc, cases, status in TRACE:
        cells = table.add_row().cells
        for i, value in enumerate((req, desc, cases, status)):
            set_cell_text(cells[i], value, size=8.3, bold=(i in (0, 3)), color=(GREEN if i == 3 else "000000"))
        set_cell_shading(cells[3], PALE_GREEN)
    set_table_geometry(table, [1500, 4200, 2400, 1260])

    doc.add_heading("Cobertura cuantitativa", level=2)
    add_kv_table(doc, [
        ("Cobertura de ejecución", "100% = 14 ejecutados / 14 planificados en el alcance seleccionado."),
        ("Tasa de aprobación", "100% = 14 aprobados / 14 ejecutados."),
        ("Cobertura de automatización", "100% = 14 automatizados / 14 seleccionados."),
        ("Cobertura de requisitos seleccionados", "100% = 12 reglas trazadas / 12 reglas definidas para esta suite."),
        ("Cobertura global del producto", "No calculable sin inventario completo y aprobado de requisitos/historias de usuario."),
        ("Backlog sugerido", f"12 escenarios pendientes identificados; no forman parte del denominador del alcance ejecutado."),
    ])
    doc.add_heading("Funcionalidades cubiertas", level=2)
    add_para(doc, "Agenda y disponibilidad; reservas concurrentes; cuotas; reprogramación; emergencias; inicio/cierre clínico; optimistic locking; estados de empresa, rol, empleado y servicio; aislamiento IDOR; activación de cuenta; carga de radiografía; historia clínica; asistencia IA; persistencia y restauración de fixtures.")
    doc.add_heading("Funcionalidades pendientes", level=2)
    add_para(doc, "Tokens inválidos/expirados, ataques de entrada, manipulación de payload, archivos inválidos/grandes, recursos inexistentes/eliminados, fallos de servicio IA, filtros/paginación, exportación, CORS/CSRF, compatibilidad multinavegador y rendimiento sostenido.")

    doc.add_heading("7. Riesgos, defectos y recomendaciones", level=1)
    doc.add_heading("Riesgos detectados", level=2)
    risks = [
        ("R-01", "Desalineación Selenium/CDP con Chrome 150", "Media", "Actualizar Selenium cuando exista soporte compatible o fijar una versión de Chrome validada."),
        ("R-02", "Suite depende de datos y fechas dinámicas", "Media", "Mantener fixtures idempotentes, reloj controlado y limpieza garantizada."),
        ("R-03", "Cobertura de seguridad parcial", "Alta", "Automatizar JWT, payload tampering, XSS/SQLi e IDs inexistentes."),
        ("R-04", "Integraciones externas no sometidas a fallos", "Alta", "Simular timeout, 4xx/5xx y respuestas inválidas del servicio IA/almacenamiento."),
        ("R-05", "Ejecución solo en Chrome/Windows", "Media", "Agregar matriz Edge/Firefox y ejecución CI reproducible."),
        ("R-06", "Videos y SlowMo aumentan duración", "Baja", "Usar video completo en evidencia formal y modo rápido en CI; conservar capturas ante fallo/final."),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("ID", "Riesgo", "Nivel", "Tratamiento")):
        set_cell_text(table.rows[0].cells[i], h, size=8.5, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    for row_data in risks:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, size=8.3, bold=(i in (0, 2)), color=(RED if i == 2 and value == "Alta" else "000000"))
    set_table_geometry(table, [900, 3200, 1000, 4260])

    doc.add_heading("Posibles defectos encontrados durante estabilización", level=2)
    add_callout(doc, "Estado actual", "No quedan defectos funcionales abiertos derivados de la ejecución final. Los hallazgos siguientes fueron defectos o debilidades de automatización/entorno ya corregidos o mitigados.", fill=PALE_GREEN, color=GREEN)
    defects = [
        ("D-E2E-01", "Conflicto de reserva no se capturaba antes de desaparecer el mensaje", "Corregido"),
        ("D-E2E-02", "Caso clínico concurrente no limpiaba estado al fallar", "Corregido"),
        ("D-E2E-03", "Selector de menú de empresa incompatible con role=menuitem", "Corregido"),
        ("D-E2E-04", "Upload esperaba visibilidad de input oculto", "Corregido"),
        ("D-E2E-05", "Historial se validaba antes de finalizar carga Angular", "Corregido"),
        ("D-ENV-01", "Advertencia CDP Chrome 150 / Selenium 4.33", "Abierto no bloqueante"),
        ("D-ENV-02", "Mojibake en algunos mensajes de consola PowerShell", "Abierto cosmético"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(("ID", "Hallazgo", "Estado")):
        set_cell_text(table.rows[0].cells[i], h, size=8.5, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    for row_data in defects:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, size=8.3, bold=(i in (0, 2)), color=(GREEN if i == 2 and value == "Corregido" else "000000"))
    set_table_geometry(table, [1400, 5900, 2060])

    doc.add_heading("Recomendaciones de mejora y automatización", level=2)
    recommendations = [
        "Mantener Selenium como herramienta oficial solicitada, Page Object Model y esperas explícitas por estado observable.",
        "Ejecutar la suite rápida sin SlowMo en CI y una ejecución con videos para entregas o hitos de QA.",
        "Publicar XML JUnit, capturas, videos y logs como artefactos versionados por ejecución.",
        "Agregar tags smoke/regression/security para seleccionar subconjuntos y reducir retroalimentación.",
        "Añadir aserciones de auditoría y base de datos mediante endpoints E2E controlados, sin acoplar los tests a detalles internos innecesarios.",
        "Automatizar primero CP-E2E-P001 a P006 por su impacto en seguridad e integridad.",
        "Si se evalúan alternativas: Playwright ofrece trazas/video nativos; Cypress simplifica depuración frontend. No migrar mientras Selenium sea requisito académico y cumpla el alcance.",
    ]
    for rec in recommendations:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(rec)
        set_run_font(r, size=10.5)

    doc.add_heading("8. Backlog de pruebas pendientes", level=1)
    add_para(doc, "Los siguientes casos fueron derivados automáticamente de las categorías solicitadas. Son propuestas y supuestos: requieren confirmar requisitos, datos y respuestas esperadas antes de automatizarse.")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(("ID", "Área", "Escenario", "Prioridad", "Estado")):
        set_cell_text(table.rows[0].cells[i], h, size=8.5, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for row_data in BACKLOG:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, size=8.2, bold=(i in (0, 3, 4)), color=(AMBER if i == 4 else "000000"))
        set_cell_shading(cells[4], PALE_AMBER)
    set_table_geometry(table, [1450, 1500, 3860, 1000, 1550])
    doc.add_heading("Pruebas de rendimiento sugeridas", level=2)
    add_kv_table(doc, [
        ("Tiempo de carga", "p95 de login, agenda e historia clínica bajo umbral acordado (Supuesto inicial: ≤ 3 s en ambiente controlado)."),
        ("Tiempo de respuesta API", "p95 y p99 de disponibilidad, creación de cita y guardado clínico; definir SLO con el equipo."),
        ("Volumen", "Agenda con miles de citas, historias extensas y paginación sostenida."),
        ("Concurrencia", "50/100/250 usuarios reservando y editando; cero doble reserva y tasa de error dentro del SLO."),
        ("Memoria", "Monitorear heap del backend y consumo del navegador durante 30-60 minutos."),
        ("Procesamiento", "Carga/consulta de archivos y análisis IA con tamaños cercanos al límite."),
    ])

    doc.add_heading("9. Evidencias y cierre", level=1)
    add_kv_table(doc, [
        ("Capturas", "target/evidence/screenshots/ — una captura final por caso."),
        ("Videos", "target/evidence/videos/ — un MP4 por caso de la ejecución con -Video."),
        ("Resultados", "target/surefire-reports/ — XML y TXT por clase de prueba."),
        ("Logs", "target/server-logs/ — backend y frontend iniciados por el runner."),
        ("Datos de entrada", "target/evidence/inputs/ — radiografia-selenium.png del caso 40."),
    ])
    doc.add_heading("Índice de evidencias de la ejecución aprobada", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("Caso", "Estado", "Captura", "Video")):
        set_cell_text(table.rows[0].cells[i], h, size=8.5, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for case in CASES:
        cells = table.add_row().cells
        values = (case["id"], "Aprobado", latest_evidence(case["num"], "png"), latest_evidence(case["num"], "mp4"))
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, size=7.8, bold=(i in (0, 1)), color=(GREEN if i == 1 else "000000"))
    set_table_geometry(table, [1300, 1200, 3430, 3430])
    doc.add_heading("Cierre QA", level=2)
    add_callout(doc, "Decisión", "APROBADO para el alcance E2E seleccionado. La suite demuestra 14/14 reglas críticas sin fallos en la ejecución final. Los riesgos y casos pendientes deben incorporarse al plan de regresión futuro según prioridad.", fill=PALE_GREEN, color=GREEN)
    add_para(doc, "Fin del documento", size=9, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    # Metadata and save.
    doc.core_properties.title = "Documentación de Pruebas E2E Selenium - VargasVet"
    doc.core_properties.subject = "Plan, diseño, ejecución y cierre QA"
    doc.core_properties.author = "Equipo QA VargasVet"
    doc.core_properties.keywords = "QA, E2E, Selenium, ISTQB, ISO 29119, VargasVet"
    doc.core_properties.comments = "Generado a partir de la suite y evidencia de ejecución aprobada del 14/07/2026."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
